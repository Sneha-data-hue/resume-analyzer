<<<<<<< HEAD
import pdfplumber
import docx

def extract_text(file):
    text = ""

    if file.name.endswith(".pdf"):
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""

    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"

=======
import pdfplumber
import docx

def extract_text(file):
    text = ""

    if file.name.endswith(".pdf"):
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""

    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"

>>>>>>> 5b96f95224f9c5fe4c27254088a6331e04a3a0be
    return text